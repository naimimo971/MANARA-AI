import os
import glob
import re
import numpy as np
from dotenv import load_dotenv
from sentence_transformers import SentenceTransformer
import faiss
from openpyxl import load_workbook
from docx import Document
from bs4 import BeautifulSoup

load_dotenv()
DATA_DIR = os.getenv("DATA_DIR", "./data")
INDEX_DIR = os.getenv("INDEX_DIR", "./kb_index")
os.makedirs(INDEX_DIR, exist_ok=True)

EMB = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")

# FIXED: Use word count, not char count
def chunk_text(text: str, max_words: int = 500, overlap: int = 50):
    if not text or not text.strip():
        return []
    words = text.split()
    if len(words) <= max_words:
        return [text]
    
    chunks = []
    i = 0
    while i < len(words):
        chunk_words = words[i:i + max_words]
        chunk = " ".join(chunk_words)
        chunks.append(chunk)
        i += max_words - overlap
        if i >= len(words):
            break
    return chunks

def extract_text(file_path):
    source = os.path.basename(file_path)
    text = ""
    try:
        ext = file_path.lower().split('.')[-1]
        print(f"Processing: {source} ({ext})")

        # SKIP PDFs — use .txt only
        if ext == 'pdf':
            print("  SKIPPED: .pdf — use convert_pdfs.py first")
            return [], source

        if ext == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            print(f"  TXT chars: {len(text)}, words: {len(text.split())}")

        elif ext == 'xlsx':
            wb = load_workbook(file_path, data_only=True)
            lines = []
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(str(c) if c is not None else "" for c in row)
                    if row_text.strip():
                        lines.append(row_text)
            text = "\n".join(lines)
            print(f"  XLSX rows: {len(lines)}")

        elif ext == 'docx':
            doc = Document(file_path)
            paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            tables = [cell.text.strip() for table in doc.tables for row in table.rows for cell in row.cells if cell.text.strip()]
            text = "\n\n".join(paras + tables)
            print(f"  DOCX paras: {len(paras)}, tables: {len(tables)}")

        elif ext == 'html':
            with open(file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f, 'html.parser')
                text = soup.get_text(separator="\n")
            print(f"  HTML chars: {len(text)}")

        else:
            print(f"  SKIPPED: unsupported format")
            return [], source

        if not text.strip():
            print(f"  WARNING: empty content")
            return [], source

        chunks = chunk_text(text)
        print(f"  → {len(chunks)} chunks created")
        return chunks, source

    except Exception as e:
        print(f"  ERROR: {e}")
        return [], source

def main():
    print(f"Scanning {DATA_DIR}...")
    files = sorted(glob.glob(os.path.join(DATA_DIR, "*")))
    print(f"Found {len(files)} files\n")

    all_chunks = []
    all_sources = []

    for file_path in files:
        chunks, source = extract_text(file_path)
        all_chunks.extend(chunks)
        all_sources.extend([source] * len(chunks))

    print(f"\nTOTAL CHUNKS: {len(all_chunks)}")
    if all_chunks:
        print("Sample sources:", all_sources[:5])

    if not all_chunks:
        print("No chunks! Run convert_pdfs.py first.")
        return

    print("Embedding...")
    embs = EMB.encode(all_chunks, batch_size=32, show_progress_bar=True, normalize_embeddings=True)

    dim = embs.shape[1]
    index = faiss.IndexFlatIP(dim)
    index.add(embs)

    faiss.write_index(index, os.path.join(INDEX_DIR, "faiss.index"))
    np.save(os.path.join(INDEX_DIR, "texts.npy"), np.array(all_chunks, dtype=object))
    np.save(os.path.join(INDEX_DIR, "sources.npy"), np.array(all_sources, dtype=object))

    print(f"\nIndex saved to {INDEX_DIR}/")
    print("Build complete!")

if __name__ == "__main__":
    main()